"""Convert the 「图卫少年」申报书 markdown into a styled .docx file.

This is a purpose-built converter for the申报书 - it understands the subset of
Markdown actually used (headings, paragraphs, GFM tables, bullet lists,
fenced code blocks, bold **text** and horizontal rules).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
HEADING_FONT = "黑体"
MONO_FONT = "Consolas"


# ---------------------------------------------------------------------------
# Low-level styling helpers
# ---------------------------------------------------------------------------


def _set_run_fonts(run, *, chinese: str = CHINESE_FONT, western: str = WESTERN_FONT) -> None:
    """Apply both Western and East Asian fonts to a run."""

    run.font.name = western
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)
    rFonts.set(qn("w:eastAsia"), chinese)


def _set_cell_shading(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "808080")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _add_bold_runs(paragraph, text: str, *, size: Pt | None = None,
                   chinese: str = CHINESE_FONT, western: str = WESTERN_FONT) -> None:
    """Add runs to *paragraph* handling **bold** markup."""

    # Split on **...** preserving the matched groups
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        if size is not None:
            run.font.size = size
        _set_run_fonts(run, chinese=chinese, western=western)


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def _configure_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = Pt(12)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), WESTERN_FONT)
    rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
    rFonts.set(qn("w:eastAsia"), CHINESE_FONT)

    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)


# ---------------------------------------------------------------------------
# Block builders
# ---------------------------------------------------------------------------


def _add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6D)
    _set_run_fonts(run, chinese=HEADING_FONT, western=HEADING_FONT)


def _add_subtitle(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _set_run_fonts(run, chinese=HEADING_FONT, western=HEADING_FONT)


def _heading_style(level: int) -> tuple[Pt, RGBColor, Pt, Pt]:
    """Return (size, color, space_before, space_after) per heading level."""

    if level == 1:
        return Pt(18), RGBColor(0x1F, 0x3A, 0x6D), Pt(18), Pt(8)
    if level == 2:
        return Pt(15), RGBColor(0x1F, 0x3A, 0x6D), Pt(14), Pt(6)
    if level == 3:
        return Pt(13), RGBColor(0x2A, 0x4A, 0x8A), Pt(10), Pt(4)
    return Pt(12), RGBColor(0x2A, 0x4A, 0x8A), Pt(8), Pt(2)


def _add_heading(doc: Document, level: int, text: str) -> None:
    paragraph = doc.add_paragraph()
    size, color, before, after = _heading_style(level)
    paragraph.paragraph_format.space_before = before
    paragraph.paragraph_format.space_after = after
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = color
    _set_run_fonts(run, chinese=HEADING_FONT, western=HEADING_FONT)


def _add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)  # 2 chars indent
    paragraph.paragraph_format.space_after = Pt(4)
    _add_bold_runs(paragraph, text)


def _add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    _add_bold_runs(paragraph, text)


def _add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    _add_bold_runs(paragraph, text)


def _add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.4)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15

    # background shading via paragraph borders + light fill
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "CCCCCC")
        pBdr.append(border)
    pPr.append(pBdr)

    body = "\n".join(lines)
    for index, line in enumerate(body.split("\n")):
        if index > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.size = Pt(10)
        _set_run_fonts(run, chinese=MONO_FONT, western=MONO_FONT)


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


def _parse_table_row(line: str) -> list[str]:
    # Strip leading/trailing pipes then split
    inner = line.strip()
    if inner.startswith("|"):
        inner = inner[1:]
    if inner.endswith("|"):
        inner = inner[:-1]
    return [cell.strip() for cell in inner.split("|")]


def _add_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for col_index, cell_text in enumerate(header):
        cell = table.cell(0, col_index)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, "1F3A6D")
        _set_cell_borders(cell)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_run_fonts(run, chinese=HEADING_FONT, western=HEADING_FONT)

    # Body rows
    for row_index, row in enumerate(rows, start=1):
        zebra = row_index % 2 == 0
        for col_index, cell_text in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)
            if zebra:
                _set_cell_shading(cell, "F5F8FC")
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            _add_bold_runs(paragraph, cell_text, size=Pt(11))


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------


def _is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped or "|" not in stripped:
        return False
    inner = stripped.strip("|")
    return all(re.fullmatch(r":?-+:?", part.strip()) for part in inner.split("|"))


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    _configure_page(doc)
    _configure_default_style(doc)

    index = 0
    in_code_block = False
    code_lines: list[str] = []

    title_seen = False
    subtitle_seen = False

    while index < len(lines):
        raw_line = lines[index]
        stripped = raw_line.strip()

        # Fenced code blocks
        if stripped.startswith("```"):
            if in_code_block:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            index += 1
            continue
        if in_code_block:
            code_lines.append(raw_line)
            index += 1
            continue

        # Horizontal rule
        if stripped == "---":
            _add_horizontal_rule(doc)
            index += 1
            continue

        # Blank line
        if not stripped:
            index += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text_content = heading_match.group(2).strip()
            if level == 1 and not title_seen:
                _add_title(doc, text_content)
                title_seen = True
            elif level == 2 and not subtitle_seen and title_seen:
                _add_subtitle(doc, text_content)
                subtitle_seen = True
            else:
                _add_heading(doc, level, text_content)
            index += 1
            continue

        # Tables: detect header row followed by separator
        if stripped.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            header = _parse_table_row(stripped)
            rows: list[list[str]] = []
            index += 2  # skip header + separator
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(_parse_table_row(lines[index]))
                index += 1
            _add_table(doc, header, rows)
            continue

        # Bullet list
        bullet_match = re.match(r"^[-*+]\s+(.*)$", stripped)
        if bullet_match:
            _add_bullet(doc, bullet_match.group(1))
            index += 1
            continue

        # Numbered list
        numbered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered_match:
            _add_numbered(doc, numbered_match.group(2))
            index += 1
            continue

        # Default: paragraph
        _add_paragraph(doc, stripped)
        index += 1

    if in_code_block and code_lines:
        _add_code_block(doc, code_lines)

    doc.save(str(docx_path))


def main(argv: Iterable[str]) -> int:
    parser = argparse.ArgumentParser(description="Convert markdown to styled docx.")
    parser.add_argument("input", type=Path, help="Path to source .md file")
    parser.add_argument("output", type=Path, help="Path to output .docx file")
    args = parser.parse_args(list(argv))

    if not args.input.exists():
        print(f"input not found: {args.input}", file=sys.stderr)
        return 1

    args.output.parent.mkdir(parents=True, exist_ok=True)
    convert(args.input, args.output)
    print(f"wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv[1:]))
